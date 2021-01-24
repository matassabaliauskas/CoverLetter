# Program that creates a DOCX format Cover Letter
# By: Matas Sabaliauskas
# Email: sabaliauskas.matas@yahoo.com

# Use:
# pip install python-docx
# pip install pypdf2

# Note: 3 Tab indentations = 12 spaces = 1 docx indentation
# General Variables: Position, PersonName, PersonAddress, CompanyName, Company Address
#
# Job Variables: Industry, JobType, CompanySize, CompanyValues,
# Self attributes: Skills, Experience, PersonalSuitability

########################################################################################################################

from random import seed
from random import randint
import docx
from docx.shared import Pt

########################################################################################################################

# seed random number generator
seed(1)
# generate some integers
for _ in range(10):
	randomvalue = randint(0, 2)     #generates random variable [0/1/2]
#	print(randomvalue)

#Personal Variables
person_name = input("Please enter your name (e.g. Tom Smith): \n")
person_phone = input("Please enter your phone (e.g. +447520123456): \n")
person_email = input("Please enter your email (e.g. name@gmail.com): \n")
person_address1 = input("Please enter your address line 1/3 (e.g. 13 Personal House): \n")
person_address2 = input("Please enter your address line 2/3 (e.g. Edinburgh): \n")
person_address3 = input("Please enter your address line 3/3 (e.g. United Kingdom): \n")
person_university = input("Please enter your university (e.g. University of Edinburgh): \n")
person_year = input("Please enter your year (e.g. 'first year'): \n")
person_degree = input("Please enter your degree (e.g. Mechanical Engineering BEng): \n")
person_interests = input("Please enter 3 things you are interested in (e.g. software, digital electronics, and hardware testing) : \n")
person_strengths = input("Please enter your 3 strengths: (e.g. hardworking, communicative, and organized): \n")

#######################################################################################################################

#Job Variables
job_title = input("Please enter internship title (e.g. Mechanical Graduate Engineer): \n")
job_find = input("Please enter where you found the opportunity (e.g. LinkedIn advertisement): \n")
company_name = input("Please enter company name (e.g. Microsoft): \n")
company_address1 = input("Please enter company address (line 1/3): (e.g. 66 Company House). Otherwise enter 0: \n")
if(company_address1 != "0"):
    company_address2 = input("Please enter company address (line 2/3): (e.g. London): \n")
    company_address3 = input("Please enter company address (line 3/3): (e.g. United Kingdom): \n")

company_person = input("Please enter the name of addressee. (If you don't know, enter 0)(e.g. Mr. John Smith): \n")
company_size = int(input("Please enter company size (1=small, 2=medium, 3=large): \n"))
company_industry = input("Please enter the industry of the company (e.g. biomedicine): \n")


####################################################################################################################

#Creating a word document
doc = docx.Document()

#Position of application
Paragraph1 = doc.add_paragraph(job_title)

#Address of person is added if company address is also inputted
if(company_address1 != "0"):
    Paragraph2 = doc.add_paragraph(person_address1 + '\n' + person_address2+ '\n' + person_address3)
    Paragraph2.alignment = 2 # for left, 1 for center, 2 right, 3 justify

#Phone Number and email
Paragraph3 = doc.add_paragraph('Phone Number: ' + person_phone + '\n' + 'Email: '+ person_email)
Paragraph3.alignment = 2 # for left, 1 for center, 2 right, 3 justify


#General Introduction
#Company Address is added if it is given
Paragraph4 = doc.add_paragraph('To Whom It May Concern' + '\n' + company_name )
if(company_address1 != "0"):
    Paragraph5 = doc.add_paragraph(company_address1 + '\n' + company_address2+ '\n' + company_address3 + '\n')

if(company_person != "0"):
    Paragraph6 = doc.add_paragraph('Dear ' + company_person)
else:
    Paragraph6 = doc.add_paragraph('Dear Sir/Madam')


#Application for a position of
Paragraph7 = doc.add_paragraph('Application for a position of ' + job_title)


###############################################################################################################


#Introduction
Paragraph8 = doc.add_paragraph('I am an ambitious student from ' + person_university + ', currently studying '
                               + person_year + ' of ' + person_degree +
                               ' degree. I am passionate about '
                               + person_interests +
                               'and I very keen to develop skills in these areas. '
                               'I have found this opportunity on '
                               +job_find+
                               ', and I am excited to gain more experience at your company.')

#Reasons for applying to a specific company
Paragraph9 = doc.add_paragraph('One of the main reasons why I would like to work at '
                               +company_name+
                               ' is high-quality, innovative, and reliable products. '
                               'I am eager to learn about both hardware and software systems, as well as find innovative solutions on their development. '
                               'I have an interest in product development, and '
                               'I have experience working with both hardware and software machines, thus I find combinations of both to be challenging and fun. '
                               'I believe that work at '
                               +company_name+
                               ' would be ideal to make significant impact through contribution in both design and technology solutions. ')


if (company_size == 1):
    # Small company comment
    Paragraph9.add_run('Working at a small company like yours will be a great way to learn about business and its development, as well as work closely with other colleagues.')
elif (company_size == 2):
    # Medium company comment
    Paragraph9.add_run('I am eager to learn from other colleagues and also contribute to the company.')
elif (company_size == 3):
    # Large company comment
    Paragraph9.add_run('One of the greatest advantages at a company of your size, would be a chance to experience workplace with a wide range of professionals, and learn from them.')


Paragraph9.add_run(' Furthermore, it would be a great opportunity to learn a range of skills, gain knowledge about product manufacturing, '
                   'as well as gain practical experience within the field of ' +company_industry+ '. ')

#Reasons why I am suitable for a job
Paragraph10 = doc.add_paragraph('I believe I would be a valuable addition to your team because I love problem solving, '
                               'I want to make an impact, and I find work with both electronics and coding genuinely fascinating. '
                               'I am also driven, creative, and passionate to learn through practical work. '
                               'I understand that the research and development projects will be challenging, '
                               'but rather than get overwhelmed I tend to tackle problems systematically. '
                               'I always try to apply my academic knowledge, solve problems through collaboration, and I generally put all of my effort and focus on the project I am working on. '
                               'Moreover, I have work-experience in different engineering firms, as well as full-time and part-time jobs while studying at the university, '
                               'as detailed further in my resume, which helped me develop both academic and practical skills.')

#Closing Paragraph
Paragraph11 = doc.add_paragraph('I have attached my CV which contains more information regarding my qualifications, experience, and skills. '
                                'If you require any documents or references, please let me know.')


if(company_person != "0"):
    Paragraphend = doc.add_paragraph('Thank you for your consideration.' + '\n\n' + 'Yours Sincerely') # When addressee is known
else:
    Paragraphend = doc.add_paragraph('Thank you for your consideration.' + '\n\n' + 'Yours Faithfully') #when we don't know the name

Paragraph10 = doc.add_paragraph(person_name)



###########################################################################################################
#Style (font/size) formating
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
Paragraph1.style = doc.styles['Normal']


#changing the page margins
# 1 inch = 914400
# top and bottom margins = 914 400 -> 457200
# side margins = 1 143 000 -> 914 400
sections = doc.sections
for section in sections:
    section.top_margin = 457200
    section.bottom_margin = 457200
    section.left_margin = 914400
    section.right_margin = 914400

#for section in sections:
#    print(section.left_margin)

#Saving the document
doc.save('CoverLetter.docx')


#################################################################################
# generator example
#if(randomvalue == 2):
#    Paragraphend = doc.add_paragraph('Thank you for your consideration.' + '\n\n' + 'Kind Regards,')

#paraobject = doc.add_paragraph('I am applying to this company because')
#paraobject.add_run('Continuing paragraph')
